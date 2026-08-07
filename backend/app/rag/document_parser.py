"""
ProspectusIQ — Multi-Format Document Parser
============================================
File: backend/app/rag/document_parser.py

Supports parsing and chunking multiple document formats for audit & analysis:
  - PDF (.pdf) via PyMuPDF (fitz)
  - Word (.docx) via python-docx
  - Excel (.xlsx, .xls, .csv, .tsv) via pandas / openpyxl
  - Plain text & Markdown (.txt, .md)
  - JSON (.json)
  - HTML (.html, .htm)

Returns normalized chunk dictionaries ready for FAISS vector indexing & compliance audit.
"""

import re
import json
import logging
from pathlib import Path
from typing import Any

log = logging.getLogger("document_parser")

SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF Document",
    ".docx": "Word Document",
    ".doc": "Word Document",
    ".xlsx": "Excel Spreadsheet",
    ".xls": "Excel Spreadsheet",
    ".csv": "CSV Spreadsheet",
    ".tsv": "TSV Spreadsheet",
    ".txt": "Plain Text",
    ".md": "Markdown Document",
    ".json": "JSON Data",
    ".html": "HTML Web Document",
    ".htm": "HTML Web Document",
}

SECTION_PATTERNS = {
    "RISK FACTORS": r"(?:RISK\s+FACTORS|SECTION\s+[IVXLC]+\s*:\s*RISK\s+FACTORS)",
    "OBJECTS OF THE ISSUE": r"(?:OBJECTS\s+OF\s+THE\s+ISSUE|PURPOSE\s+OF\s+THE\s+ISSUE)",
    "BUSINESS OVERVIEW": r"(?:OUR\s+BUSINESS|BUSINESS\s+OVERVIEW|ABOUT\s+OUR\s+COMPANY)",
    "FINANCIAL INFORMATION": r"(?:FINANCIAL\s+INFORMATION|FINANCIAL\s+STATEMENTS|RESTATED\s+FINANCIAL|BALANCE\s+SHEET)",
    "PROMOTERS AND PROMOTER GROUP": r"(?:OUR\s+PROMOTERS|PROMOTER\s+GROUP)",
    "PEER GROUP COMPARISON": r"(?:PEER\s+GROUP|COMPARISON\s+WITH\s+LISTED\s+PEERS)",
    "GENERAL INFORMATION": r"(?:GENERAL\s+INFORMATION|DEFINITIONS?\s+AND\s+ABBREVIATIONS?)",
    "CAPITAL STRUCTURE": r"(?:CAPITAL\s+STRUCTURE)",
    "DIVIDEND POLICY": r"(?:DIVIDEND\s+POLICY)",
    "INDUSTRY OVERVIEW": r"(?:INDUSTRY\s+OVERVIEW|OUR\s+INDUSTRY)",
    "LEGAL AND OTHER INFORMATION": r"(?:LEGAL\s+AND\s+OTHER|OUTSTANDING\s+LITIGATION|GOVERNMENT\s+APPROVALS)",
}


def clean_text(raw_text: str) -> str:
    """Normalize text by stripping null bytes, HTML tags, and collapsing space."""
    if not raw_text:
        return ""
    text = raw_text.replace("\x00", "").replace("\ufffd", "")
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def text_to_chunks(text: str, source_name: str, doc_type: str, section: str = "GENERAL", page_num: int = 1, chunk_size: int = 350) -> list[dict[str, Any]]:
    """Chunk clean text into 250-400 character blocks with metadata."""
    chunks = []
    lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 15]
    buf = []
    
    current_sec = section
    for line in lines:
        for sec_name, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, line, re.IGNORECASE):
                current_sec = sec_name
                break
                
        buf.append(line)
        if sum(len(x) for x in buf) >= chunk_size:
            joined = " ".join(buf)
            chunks.append({
                "chunk_id": f"{Path(source_name).stem}_p{page_num}_c{len(chunks) + 1}",
                "pdf_source": source_name,
                "page": page_num,
                "section": current_sec,
                "doc_type": doc_type,
                "text": joined,
            })
            buf = []

    if buf:
        joined = " ".join(buf)
        if len(joined) > 20:
            chunks.append({
                "chunk_id": f"{Path(source_name).stem}_p{page_num}_c{len(chunks) + 1}",
                "pdf_source": source_name,
                "page": page_num,
                "section": current_sec,
                "doc_type": doc_type,
                "text": joined,
            })

    return chunks


# ── Format-Specific Parsers ───────────────────────────────────────────────────

def parse_pdf(file_path: Path, doc_type: str) -> list[dict[str, Any]]:
    """Parse PDF document (.pdf) using PyMuPDF."""
    import fitz
    chunks = []
    try:
        doc = fitz.open(str(file_path))
        for page_idx in range(len(doc)):
            try:
                page = doc.load_page(page_idx)
                raw_text = page.get_text("text")
                cleaned = clean_text(raw_text)
                if cleaned:
                    chunks.extend(text_to_chunks(cleaned, file_path.name, doc_type, page_num=page_idx + 1))
            except Exception:
                continue
        doc.close()
    except Exception as e:
        log.error(f"Error parsing PDF {file_path.name}: {e}")
    return chunks


def parse_docx(file_path: Path, doc_type: str) -> list[dict[str, Any]]:
    """Parse Word document (.docx) using python-docx."""
    import docx
    chunks = []
    try:
        doc = docx.Document(str(file_path))
        full_text = []
        
        for para in doc.paragraphs:
            txt = clean_text(para.text)
            if txt:
                full_text.append(txt)
                
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_txt = [clean_text(cell.text) for cell in row.cells if cell.text.strip()]
                if row_txt:
                    full_text.append(" | ".join(row_txt))
                    
        joined_text = "\n".join(full_text)
        chunks = text_to_chunks(joined_text, file_path.name, doc_type)
    except Exception as e:
        log.error(f"Error parsing Word DOCX {file_path.name}: {e}")
    return chunks


def parse_excel_or_csv(file_path: Path, doc_type: str) -> list[dict[str, Any]]:
    """Parse Excel spreadsheet (.xlsx, .csv, .tsv) using pandas."""
    import pandas as pd
    chunks = []
    try:
        ext = file_path.suffix.lower()
        dfs: dict[str, pd.DataFrame] = {}
        
        if ext in [".csv", ".tsv"]:
            sep = "\t" if ext == ".tsv" else ","
            dfs["Sheet1"] = pd.read_csv(str(file_path), sep=sep)
        else:
            dfs = pd.read_excel(str(file_path), sheet_name=None)
            
        for sheet_name, df in dfs.items():
            if df.empty:
                continue
            df_clean = df.fillna("")
            records = df_clean.to_dict(orient="records")
            sheet_text_blocks = []
            
            for idx, row in enumerate(records, 1):
                row_str = ", ".join(f"{k}: {v}" for k, v in row.items() if str(v).strip() != "")
                if row_str:
                    sheet_text_blocks.append(f"Row {idx} -> {row_str}")
                    
            joined = f"Sheet: {sheet_name}\n" + "\n".join(sheet_text_blocks)
            chunks.extend(text_to_chunks(joined, file_path.name, doc_type, section=f"EXCEL_{sheet_name.upper()}"))
    except Exception as e:
        log.error(f"Error parsing Excel/CSV {file_path.name}: {e}")
    return chunks


def parse_plain_text(file_path: Path, doc_type: str) -> list[dict[str, Any]]:
    """Parse plain text, Markdown, or HTML files (.txt, .md, .html, .json)."""
    chunks = []
    try:
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        if file_path.suffix.lower() == ".json":
            try:
                obj = json.loads(raw)
                raw = json.dumps(obj, indent=2)
            except Exception:
                pass
        cleaned = clean_text(raw)
        chunks = text_to_chunks(cleaned, file_path.name, doc_type)
    except Exception as e:
        log.error(f"Error parsing text/markdown {file_path.name}: {e}")
    return chunks


# ── Universal Parser Interface ────────────────────────────────────────────────

def parse_document(file_path: Path | str, doc_type: str = "General") -> list[dict[str, Any]]:
    """
    Universal document parser for ProspectusIQ.
    Automatically detects format by extension and extracts structured chunks.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if not path.exists():
        log.error(f"File not found: {path}")
        return []

    log.info(f"Parsing document [{SUPPORTED_EXTENSIONS.get(ext, 'Unknown')}]: {path.name}")

    if ext == ".pdf":
        return parse_pdf(path, doc_type)
    elif ext in [".docx", ".doc"]:
        return parse_docx(path, doc_type)
    elif ext in [".xlsx", ".xls", ".csv", ".tsv"]:
        return parse_excel_or_csv(path, doc_type)
    elif ext in [".txt", ".md", ".json", ".html", ".htm"]:
        return parse_plain_text(path, doc_type)
    else:
        log.warning(f"Unsupported extension '{ext}' for {path.name}, attempting text fallback...")
        return parse_plain_text(path, doc_type)
